"""重新產生「技能清單.docx」——彙整 `.claude/skills/*/SKILL.md` 的技能清單、model 分級、觸發詞、
用途摘要。每次新增/修改技能後手動重跑一次即可同步文件，不是排程腳本。
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SKILLS_DIR = Path(__file__).parent / ".claude" / "skills"
OUTPUT_PATH = Path(__file__).parent / "技能清單.docx"

MODEL_LABEL = {
    "haiku": "haiku（機械式操作、不太會出錯）",
    "sonnet": "sonnet（需要判斷/摘要的中等複雜度）",
    "opus": "opus（需要診斷推理、誤判成本高）",
}


def _parse_skill(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    front_match = re.match(r"^---\n(.*?)\n---\n(.*)$", text, re.DOTALL)
    front, body = front_match.group(1), front_match.group(2)

    fields = {}
    for line in front.splitlines():
        if ":" in line:
            key, _, value = line.partition(":")
            fields[key.strip()] = value.strip()

    title_match = re.search(r"^#\s+(.+)$", body, re.MULTILINE)
    title = title_match.group(1) if title_match else fields.get("name", "")

    return {
        "name": fields.get("name", path.parent.name),
        "model": fields.get("model", "inherit"),
        "description": fields.get("description", ""),
        "title": title,
    }


def build_document() -> Document:
    skills = sorted(
        (_parse_skill(p) for p in SKILLS_DIR.glob("*/SKILL.md")),
        key=lambda s: s["name"],
    )

    doc = Document()

    title = doc.add_heading("AI 投資顧問專案 — Claude Code 技能清單", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "本文件彙整專案 `.claude/skills/` 下的所有 Claude Code 技能，"
        "包含觸發詞、使用的模型分級、用途摘要。model 分級規則見專案 `CLAUDE.md`："
        "haiku 用於機械式操作、不太會出錯的簡單任務；sonnet 用於需要判斷/摘要的中等複雜度任務；"
        "opus 用於需要診斷推理、誤判成本高的複雜任務。"
    )
    intro.runs[0].font.size = Pt(10)

    doc.add_paragraph(f"共 {len(skills)} 個技能。", style="Intense Quote")

    for skill in skills:
        doc.add_heading(f"/{skill['name']}", level=1)

        p = doc.add_paragraph()
        p.add_run("使用模型：").bold = True
        p.add_run(MODEL_LABEL.get(skill["model"], skill["model"]))

        p = doc.add_paragraph()
        p.add_run("觸發詞與用途：").bold = True
        doc.add_paragraph(skill["description"])

    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"已產生 {OUTPUT_PATH}")
